import io
import os
import re
import requests
import docx
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS

app = Flask(__name__)
CORS(app)

def extract_text_from_docx(docx_stream):
    """Extracts all text from a docx file."""
    doc = docx.Document(docx_stream)
    
    # Extract text from paragraphs and tables
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)

    # Join the extracted text
    raw_text = '\n'.join(full_text)

    # Clean up the text by replacing multiple newlines or spaces with a single space
    cleaned_text = re.sub(r'\s+', ' ', raw_text)
    
    # Remove spaces before punctuation
    cleaned_text = re.sub(r'\s([?.!,"](?:\s|$))', r'\1', cleaned_text)
    
    # Remove leading and trailing whitespace
    cleaned_text = cleaned_text.strip()

    return cleaned_text

def find_replace_text(doc, search_strings, replace_strings):
    """Find and replace text in the docx document."""
    for paragraph in doc.paragraphs:
        for i, search_string in enumerate(search_strings):
            if search_string in paragraph.text:
                paragraph.text = paragraph.text.replace(search_string, replace_strings[i])

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for i, search_string in enumerate(search_strings):
                        if search_string in paragraph.text:
                            paragraph.text = paragraph.text.replace(search_string, replace_strings[i])

@app.route('/extract_text', methods=['POST'])
def extract_text():
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    if file:
        # Read the file into a bytes stream
        file_stream = io.BytesIO(file.read())
        
        # Extract text from the document
        extracted_text = extract_text_from_docx(file_stream)

        return jsonify({
            "extracted_text": extracted_text
        })

@app.route('/job', methods=['POST'])
def job():
    if 'file_url' not in request.json:
        return jsonify({"error": "File URL is required"}), 400

    file_url = request.json['file_url']

    if 'searchStrings' not in request.json or 'replaceStrings' not in request.json:
        return jsonify({"error": "Search strings and replace strings are required"}), 400

    search_strings = request.json['searchStrings']
    replace_strings = request.json['replaceStrings']

    if len(search_strings) != len(replace_strings):
        return jsonify({"error": "Number of search strings and replace strings must match"}), 400

    # Download the file from the provided URL with a timeout
    try:
        response = requests.get(file_url, timeout=10)  # Set a 10-second timeout
        response.raise_for_status()  # Check if the download was successful
    except requests.exceptions.Timeout:
        return jsonify({"error": "The request timed out"}), 500
    except requests.exceptions.RequestException as e:
        return jsonify({"error": f"Failed to download the file: {e}"}), 500

    file_stream = io.BytesIO(response.content)

    # Perform replacements in the document
    doc = docx.Document(file_stream)
    find_replace_text(doc, search_strings, replace_strings)

    # Save the updated document to a BytesIO object
    output_stream = io.BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)  # Rewind the stream to the beginning

    # Send the modified file as a response
    return send_file(
        output_stream,
        as_attachment=True,
        download_name="updated_document.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == '__main__':
    app.run(debug=True)
